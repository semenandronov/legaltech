"""File parsing utilities for Legal AI Vault"""
import os
from typing import Tuple
from pypdf import PdfReader
from docx import Document
import openpyxl


def _clean_text(text: str) -> str:
    """Remove NUL bytes and BOM markers"""
    if text is None:
        return ""
    return text.replace("\x00", "").replace("\ufeff", "")


def parse_pdf(file_content: bytes, filename: str) -> str:
    """Extract text from PDF file"""
    try:
        import io
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text())
        return _clean_text("\n".join(text_parts))
    except Exception as e:
        raise ValueError(f"Ошибка при чтении PDF файла {filename}: {str(e)}")


def parse_docx(file_content: bytes, filename: str) -> str:
    """Extract text from DOCX file"""
    try:
        import io
        import zipfile
        
        # Validate that file is actually a ZIP (DOCX files are ZIP archives)
        try:
            zip_test = zipfile.ZipFile(io.BytesIO(file_content))
            zip_test.close()
        except zipfile.BadZipFile:
            raise ValueError(f"Файл {filename} не является корректным DOCX файлом (не является ZIP архивом)")
        
        # Parse DOCX
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        result = _clean_text("\n".join(text_parts))
        if not result.strip():
            raise ValueError(f"DOCX файл {filename} не содержит текста")
        return result
    except ValueError:
        raise
    except Exception as e:
        raise ValueError(f"Ошибка при чтении DOCX файла {filename}: {str(e)}")


def parse_txt(file_content: bytes, filename: str) -> str:
    """Extract text from TXT file"""
    try:
        # Try different encodings
        for encoding in ['utf-8-sig', 'utf-8', 'cp1251', 'latin-1']:
            try:
                return _clean_text(file_content.decode(encoding))
            except UnicodeDecodeError:
                continue
        raise ValueError(f"Не удалось декодировать файл {filename}")
    except Exception as e:
        raise ValueError(f"Ошибка при чтении TXT файла {filename}: {str(e)}")


def parse_xlsx(file_content: bytes, filename: str) -> str:
    """Extract text from XLSX file"""
    try:
        import io
        xlsx_file = io.BytesIO(file_content)
        workbook = openpyxl.load_workbook(xlsx_file, data_only=True)
        text_parts = []
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text_parts.append(f"[Лист: {sheet_name}]")
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    text_parts.append(row_text)
        return _clean_text("\n".join(text_parts))
    except Exception as e:
        raise ValueError(f"Ошибка при чтении XLSX файла {filename}: {str(e)}")


def parse_file(file_content: bytes, filename: str) -> str:
    """Parse file based on extension"""
    _, ext = os.path.splitext(filename.lower())
    
    if ext == ".pdf":
        return parse_pdf(file_content, filename)
    elif ext == ".docx":
        return parse_docx(file_content, filename)
    elif ext == ".txt":
        return parse_txt(file_content, filename)
    elif ext == ".xlsx":
        return parse_xlsx(file_content, filename)
    else:
        raise ValueError(f"Неподдерживаемый формат файла: {ext}")

