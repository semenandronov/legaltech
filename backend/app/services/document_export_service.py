"""Document Export Service for Legal AI Vault"""
from typing import Optional
from sqlalchemy.orm import Session
from app.models.document_editor import Document
from app.services.document_editor_service import DocumentEditorService
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table as ReportLabTable, TableStyle
from reportlab.lib import colors
from io import BytesIO
import logging
import re
from html.parser import HTMLParser
from html import unescape

logger = logging.getLogger(__name__)


class HTMLParserForExport(HTMLParser):
    """Parser for converting HTML to structured content"""
    def __init__(self):
        super().__init__()
        self.elements = []
        self.current_tag = None
        self.current_text = []
        self.current_attrs = {}
    
    def handle_starttag(self, tag, attrs):
        # Save any accumulated text
        if self.current_text:
            text = ' '.join(self.current_text).strip()
            if text:
                self.elements.append({
                    'type': 'text',
                    'tag': self.current_tag or 'p',
                    'text': text,
                    'attrs': self.current_attrs
                })
            self.current_text = []
        
        self.current_tag = tag
        self.current_attrs = dict(attrs)
    
    def handle_endtag(self, tag):
        # Save any accumulated text
        if self.current_text:
            text = ' '.join(self.current_text).strip()
            if text:
                self.elements.append({
                    'type': 'text',
                    'tag': self.current_tag or 'p',
                    'text': text,
                    'attrs': self.current_attrs
                })
            self.current_text = []
        
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'td', 'th']:
            self.elements.append({
                'type': 'end',
                'tag': tag
            })
        
        self.current_tag = None
        self.current_attrs = {}
    
    def handle_data(self, data):
        self.current_text.append(data.strip())
    
    def get_elements(self):
        # Process any remaining text
        if self.current_text:
            text = ' '.join(self.current_text).strip()
            if text:
                self.elements.append({
                    'type': 'text',
                    'tag': self.current_tag or 'p',
                    'text': text,
                    'attrs': self.current_attrs
                })
        return self.elements


class DocumentExportService:
    """Service for exporting documents to various formats"""
    
    def __init__(self, db: Session):
        """Initialize document export service"""
        self.db = db
        self.document_service = DocumentEditorService(db)
    
    def export_to_docx(
        self,
        document_id: str,
        user_id: str
    ) -> BytesIO:
        """
        Export document to DOCX format
        
        Args:
            document_id: Document identifier
            user_id: User identifier
            
        Returns:
            BytesIO buffer with DOCX file
        """
        document = self.document_service.get_document(document_id, user_id)
        if not document:
            raise ValueError(f"Document {document_id} not found")
        
        # Create Word document
        doc = DocxDocument()
        
        # Add title
        title = doc.add_heading(document.title, 0)
        
        # Parse HTML content
        html_content = document.content
        if not html_content:
            html_content = "<p>Пустой документ</p>"
        
        # Simple HTML parsing and conversion
        # Remove script and style tags
        html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Parse HTML
        parser = HTMLParserForExport()
        parser.feed(html_content)
        elements = parser.get_elements()
        
        # Convert elements to DOCX
        for element in elements:
            if element['type'] == 'text':
                tag = element['tag']
                text = unescape(element['text'])
                
                if not text:
                    continue
                
                # Handle different HTML tags
                if tag == 'h1':
                    doc.add_heading(text, level=1)
                elif tag == 'h2':
                    doc.add_heading(text, level=2)
                elif tag == 'h3':
                    doc.add_heading(text, level=3)
                elif tag == 'h4':
                    doc.add_heading(text, level=4)
                elif tag == 'h5':
                    doc.add_heading(text, level=5)
                elif tag == 'h6':
                    doc.add_heading(text, level=6)
                elif tag == 'p':
                    # Check for formatting
                    attrs = element.get('attrs', {})
                    para = doc.add_paragraph(text)
                    
                    # Apply formatting if needed
                    if attrs.get('style'):
                        # Simple style parsing (can be enhanced)
                        if 'font-weight:bold' in attrs['style'] or 'font-weight: bold' in attrs['style']:
                            for run in para.runs:
                                run.bold = True
                        if 'font-style:italic' in attrs['style'] or 'font-style: italic' in attrs['style']:
                            for run in para.runs:
                                run.italic = True
                elif tag == 'li':
                    # List item
                    para = doc.add_paragraph(text, style='List Bullet')
                elif tag in ['strong', 'b']:
                    para = doc.add_paragraph(text)
                    for run in para.runs:
                        run.bold = True
                elif tag in ['em', 'i']:
                    para = doc.add_paragraph(text)
                    for run in para.runs:
                        run.italic = True
                else:
                    # Default to paragraph
                    doc.add_paragraph(text)
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info(f"Exported document {document_id} to DOCX")
        return buffer
    
    def export_to_pdf(
        self,
        document_id: str,
        user_id: str
    ) -> BytesIO:
        """
        Export document to PDF format
        
        Args:
            document_id: Document identifier
            user_id: User identifier
            
        Returns:
            BytesIO buffer with PDF file
        """
        document = self.document_service.get_document(document_id, user_id)
        if not document:
            raise ValueError(f"Document {document_id} not found")
        
        # Create PDF document
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        styles = getSampleStyleSheet()
        
        # Add title
        title_style = ParagraphStyle(
            'DocumentTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=12
        )
        story.append(Paragraph(document.title, title_style))
        story.append(Spacer(1, 12))
        
        # Parse HTML content
        html_content = document.content
        if not html_content:
            html_content = "<p>Пустой документ</p>"
        
        # Remove script and style tags
        html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        html_content = re.sub(r'<style[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
        
        # Parse HTML
        parser = HTMLParserForExport()
        parser.feed(html_content)
        elements = parser.get_elements()
        
        # Convert elements to PDF
        for element in elements:
            if element['type'] == 'text':
                tag = element['tag']
                text = unescape(element['text'])
                
                if not text:
                    continue
                
                # Handle different HTML tags
                if tag == 'h1':
                    story.append(Paragraph(text, styles['Heading1']))
                    story.append(Spacer(1, 12))
                elif tag == 'h2':
                    story.append(Paragraph(text, styles['Heading2']))
                    story.append(Spacer(1, 10))
                elif tag == 'h3':
                    story.append(Paragraph(text, styles['Heading3']))
                    story.append(Spacer(1, 8))
                elif tag == 'h4':
                    story.append(Paragraph(text, styles['Heading4']))
                    story.append(Spacer(1, 6))
                elif tag == 'h5':
                    story.append(Paragraph(text, styles['Heading5']))
                    story.append(Spacer(1, 4))
                elif tag == 'h6':
                    story.append(Paragraph(text, styles['Heading6']))
                    story.append(Spacer(1, 4))
                elif tag == 'p':
                    story.append(Paragraph(text, styles['Normal']))
                    story.append(Spacer(1, 6))
                elif tag == 'li':
                    story.append(Paragraph(f"• {text}", styles['Normal']))
                    story.append(Spacer(1, 4))
                elif tag in ['strong', 'b']:
                    # Bold text
                    bold_style = ParagraphStyle(
                        'BoldText',
                        parent=styles['Normal'],
                        fontName='Helvetica-Bold'
                    )
                    story.append(Paragraph(text, bold_style))
                    story.append(Spacer(1, 6))
                elif tag in ['em', 'i']:
                    # Italic text
                    italic_style = ParagraphStyle(
                        'ItalicText',
                        parent=styles['Normal'],
                        fontName='Helvetica-Oblique'
                    )
                    story.append(Paragraph(text, italic_style))
                    story.append(Spacer(1, 6))
                else:
                    # Default to normal paragraph
                    story.append(Paragraph(text, styles['Normal']))
                    story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        logger.info(f"Exported document {document_id} to PDF")
        return buffer
















