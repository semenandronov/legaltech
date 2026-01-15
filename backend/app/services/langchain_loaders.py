"""LangChain document loaders for Legal AI Vault"""
from typing import List, Dict, Any, Optional
from langchain_community.document_loaders import (
    PyPDFLoader,
    UnstructuredWordDocumentLoader,
    TextLoader,
    CSVLoader,
)
from langchain_core.documents import Document
from app.config import config
import logging
import io
import tempfile
import os

logger = logging.getLogger(__name__)


class DocumentLoaderService:
    """Service for loading documents using LangChain loaders"""
    
    @staticmethod
    def load_pdf(content: bytes, filename: str) -> List[Document]:
        """
        Load PDF document using PyPDFLoader
        
        Args:
            content: PDF file content as bytes
            filename: Original filename
            
        Returns:
            List of Document objects with metadata
        """
        try:
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(content)
                tmp_path = tmp_file.name
            
            try:
                loader = PyPDFLoader(tmp_path)
                documents = loader.load()
                
                # Add filename and page info to metadata
                for doc in documents:
                    # Устанавливаем source_file для всех документов
                    doc.metadata["source_file"] = filename
                    
                    # Извлекаем номер страницы из source метаданных PyPDFLoader
                    source = doc.metadata.get("source", "")
                    if source and "page" in source.lower():
                        try:
                            # PyPDFLoader формат: ".../filename.pdf:page=N"
                            # Ищем номер страницы
                            import re
                            page_match = re.search(r'page[=:]?\s*(\d+)', source, re.IGNORECASE)
                            if page_match:
                                page_num = int(page_match.group(1))
                                doc.metadata["source_page"] = page_num
                        except Exception as e:
                            logger.debug(f"Could not extract page number from source '{source}': {e}")
                
                logger.info(f"Loaded PDF {filename}: {len(documents)} pages")
                return documents
            finally:
                # Clean up temp file
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
        except Exception as e:
            logger.error(f"Error loading PDF {filename}: {e}", exc_info=True)
            raise
    
    @staticmethod
    def load_docx(content: bytes, filename: str) -> List[Document]:
        """
        Load DOCX document using UnstructuredWordDocumentLoader
        
        Args:
            content: DOCX file content as bytes
            filename: Original filename
            
        Returns:
            List of Document objects with metadata
        """
        # Проверяем, что файл похож на DOCX (ZIP архив с определённой структурой)
        if len(content) < 100:
            raise ValueError(f"Файл {filename} слишком маленький для DOCX формата")
        
        # DOCX - это ZIP архив, должен начинаться с PK (0x50, 0x4B)
        if content[:2] != b'PK':
            logger.warning(f"File {filename} does not have DOCX signature (PK header), might not be a valid DOCX")
            # Пытаемся обработать как текстовый файл
            try:
                text_content = content.decode('utf-8', errors='ignore')
                if text_content.strip():
                    logger.info(f"Loaded {filename} as plain text (not a valid DOCX)")
                    return [Document(page_content=text_content, metadata={"source_file": filename, "note": "loaded as text"})]
            except Exception:
                pass
            raise ValueError(
                f"Файл {filename} не является корректным DOCX файлом. "
                f"Убедитесь, что вы загружаете правильный файл."
            )
        
        try:
            # Save to temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                tmp_file.write(content)
                tmp_path = tmp_file.name
            
            try:
                documents = None
                unstructured_error = None
                docx_error = None
                
                # Try python-docx first (more reliable for basic DOCX)
                try:
                    from docx import Document as DocxDocument
                    doc = DocxDocument(io.BytesIO(content))
                    
                    # Извлекаем текст из параграфов
                    paragraphs_text = [para.text for para in doc.paragraphs if para.text.strip()]
                    
                    # Также извлекаем текст из таблиц
                    tables_text = []
                    for table in doc.tables:
                        for row in table.rows:
                            row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                            if row_text:
                                tables_text.append(row_text)
                    
                    all_text = "\n".join(paragraphs_text)
                    if tables_text:
                        all_text += "\n\n[Таблицы]\n" + "\n".join(tables_text)
                    
                    if all_text.strip():
                        documents = [Document(page_content=all_text, metadata={"source_file": filename})]
                        logger.info(f"Loaded DOCX {filename} using python-docx: {len(all_text)} chars")
                except Exception as e:
                    docx_error = e
                    logger.debug(f"python-docx failed for {filename}: {e}")
                
                # If python-docx failed or returned empty, try UnstructuredWordDocumentLoader
                if not documents or not documents[0].page_content.strip():
                    try:
                        loader = UnstructuredWordDocumentLoader(tmp_path, mode="elements")
                        documents = loader.load()
                        logger.info(f"Loaded DOCX {filename} using UnstructuredWordDocumentLoader: {len(documents)} elements")
                    except Exception as e:
                        unstructured_error = e
                        logger.debug(f"UnstructuredWordDocumentLoader failed for {filename}: {e}")
                
                # If both methods failed
                if not documents or not any(doc.page_content.strip() for doc in documents):
                    error_msgs = []
                    if docx_error:
                        error_msgs.append(f"python-docx: {str(docx_error)[:100]}")
                    if unstructured_error:
                        error_msgs.append(f"unstructured: {str(unstructured_error)[:100]}")
                    
                    raise ValueError(
                        f"Не удалось извлечь текст из файла {filename}. "
                        f"Возможно, файл поврежден или имеет нестандартный формат. "
                        f"Ошибки: {'; '.join(error_msgs) if error_msgs else 'неизвестная ошибка'}"
                    )
                
                # Add filename to metadata
                for doc in documents:
                    if "source_file" not in doc.metadata:
                        doc.metadata["source_file"] = filename
                
                return documents
            finally:
                # Clean up temp file
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
        except ValueError:
            raise
        except Exception as e:
            logger.error(f"Error loading DOCX {filename}: {e}", exc_info=True)
            raise ValueError(
                f"Ошибка при загрузке файла {filename}: {str(e)[:200]}. "
                f"Убедитесь, что файл является корректным DOCX файлом."
            )
    
    @staticmethod
    def load_txt(content: bytes, filename: str) -> List[Document]:
        """
        Load text document using TextLoader
        
        Args:
            content: Text file content as bytes
            filename: Original filename
            
        Returns:
            List of Document objects with metadata
        """
        try:
            # Decode content
            try:
                text_content = content.decode('utf-8')
            except UnicodeDecodeError:
                text_content = content.decode('latin-1')
            
            # Save to temporary file
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.txt', encoding='utf-8') as tmp_file:
                tmp_file.write(text_content)
                tmp_path = tmp_file.name
            
            try:
                loader = TextLoader(tmp_path, encoding='utf-8')
                documents = loader.load()
                
                # Add filename to metadata
                for doc in documents:
                    doc.metadata["source_file"] = filename
                
                logger.info(f"Loaded TXT {filename}: {len(documents)} documents")
                return documents
            finally:
                # Clean up temp file
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)
        except Exception as e:
            logger.error(f"Error loading TXT {filename}: {e}", exc_info=True)
            raise
    
    @staticmethod
    def load_xlsx(content: bytes, filename: str) -> List[Document]:
        """
        Load Excel document (convert to text format)
        
        Args:
            content: XLSX file content as bytes
            filename: Original filename
            
        Returns:
            List of Document objects with metadata
        """
        try:
            from openpyxl import load_workbook
            
            # Load workbook
            workbook = load_workbook(io.BytesIO(content), data_only=True)
            
            documents = []
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Extract text from cells
                rows_text = []
                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        rows_text.append(row_text)
                
                if rows_text:
                    sheet_text = "\n".join(rows_text)
                    doc = Document(
                        page_content=sheet_text,
                        metadata={
                            "source_file": filename,
                            "sheet_name": sheet_name,
                            "source_type": "excel"
                        }
                    )
                    documents.append(doc)
            
            logger.info(f"Loaded XLSX {filename}: {len(documents)} sheets")
            return documents
        except Exception as e:
            logger.error(f"Error loading XLSX {filename}: {e}", exc_info=True)
            raise
    
    @staticmethod
    def load_document(content: bytes, filename: str) -> List[Document]:
        """
        Load document based on file extension
        
        Args:
            content: File content as bytes
            filename: Original filename
            
        Returns:
            List of Document objects with metadata
        """
        # Get file extension
        _, ext = os.path.splitext(filename)
        ext = ext.lower()
        
        if ext == '.pdf':
            return DocumentLoaderService.load_pdf(content, filename)
        elif ext in ['.docx', '.doc']:
            return DocumentLoaderService.load_docx(content, filename)
        elif ext == '.txt':
            return DocumentLoaderService.load_txt(content, filename)
        elif ext in ['.xlsx', '.xls']:
            return DocumentLoaderService.load_xlsx(content, filename)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

