"""Report generator service"""
from typing import Dict, Any, List
from sqlalchemy.orm import Session
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from io import BytesIO
import logging

logger = logging.getLogger(__name__)


class ReportGenerator:
    """Service for generating reports in various formats"""
    
    def __init__(self, db: Session):
        """Initialize report generator"""
        self.db = db
    
    def generate_executive_summary(
        self,
        case_id: str,
        summary: str,
        key_facts: Dict[str, Any],
        risk_analysis: str = None
    ) -> BytesIO:
        """
        Generate Executive Summary report (Word)
        
        Args:
            case_id: Case identifier
            summary: Summary text
            key_facts: Key facts dictionary
            risk_analysis: Risk analysis text (optional)
            
        Returns:
            BytesIO buffer with Word document
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('LEGALCHAIN AI', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('EXECUTIVE SUMMARY', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Case info
        doc.add_paragraph(f'Дело: {case_id}')
        doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
        doc.add_paragraph('')
        
        # Summary section
        doc.add_heading('СУТЬ ДЕЛА', 2)
        doc.add_paragraph(summary)
        doc.add_paragraph('')
        
        # Key facts section
        if key_facts:
            doc.add_heading('КЛЮЧЕВЫЕ ФАКТЫ', 2)
            if key_facts.get('parties'):
                doc.add_paragraph(f"Истец: {key_facts['parties'].get('plaintiff', 'Не указан')}")
                doc.add_paragraph(f"Ответчик: {key_facts['parties'].get('defendant', 'Не указан')}")
            if key_facts.get('amounts'):
                doc.add_paragraph(f"Сумма спора: {key_facts['amounts'].get('dispute_amount', 'Не указана')}")
            doc.add_paragraph('')
        
        # Risk analysis section
        if risk_analysis:
            doc.add_heading('РИСК-ОЦЕНКА', 2)
            doc.add_paragraph(risk_analysis)
            doc.add_paragraph('')
        
        # Footer
        doc.add_paragraph('')
        doc.add_paragraph('Отчет подготовлен LEGALCHAIN AI (система анализа документов)')
        doc.add_paragraph('Используй этот отчет в информационных целях.')
        doc.add_paragraph('ПЕРЕПРОВЕРЬ все выводы перед использованием в суде!')
        
        # Save to BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def generate_detailed_analysis(
        self,
        case_id: str,
        timeline: List[Dict[str, Any]],
        discrepancies: List[Dict[str, Any]],
        key_facts: Dict[str, Any],
        summary: str,
        risk_analysis: str = None
    ) -> BytesIO:
        """
        Generate Detailed Analysis report (Word)
        
        Args:
            case_id: Case identifier
            timeline: Timeline events
            discrepancies: Discrepancies list
            key_facts: Key facts dictionary
            summary: Summary text
            risk_analysis: Risk analysis text (optional)
            
        Returns:
            BytesIO buffer with Word document
        """
        doc = Document()
        
        # Title
        title = doc.add_heading('LEGALCHAIN AI', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('DETAILED ANALYSIS', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f'Дело: {case_id}')
        doc.add_paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y")}')
        doc.add_paragraph('')
        
        # Summary
        doc.add_heading('КРАТКОЕ РЕЗЮМЕ', 2)
        doc.add_paragraph(summary)
        doc.add_paragraph('')
        
        # Timeline
        if timeline:
            doc.add_heading('ТАЙМЛАЙН', 2)
            for event in timeline:
                event_date = event.get('date', '')
                event_desc = event.get('description', '')
                event_source = event.get('source_document', '')
                doc.add_paragraph(f"{event_date}: {event_desc} (Источник: {event_source})")
            doc.add_paragraph('')
        
        # Discrepancies
        if discrepancies:
            doc.add_heading('ПРОТИВОРЕЧИЯ', 2)
            for i, disc in enumerate(discrepancies, 1):
                doc.add_paragraph(f"{i}. {disc.get('type', 'Неизвестный тип')} - {disc.get('severity', 'MEDIUM')}")
                doc.add_paragraph(f"   {disc.get('description', '')}")
                if disc.get('source_documents'):
                    doc.add_paragraph(f"   Источники: {', '.join(disc['source_documents'])}")
            doc.add_paragraph('')
        
        # Key facts
        if key_facts:
            doc.add_heading('КЛЮЧЕВЫЕ ФАКТЫ', 2)
            for key, value in key_facts.items():
                if isinstance(value, dict):
                    doc.add_paragraph(f"{key}:")
                    for k, v in value.items():
                        doc.add_paragraph(f"  - {k}: {v}")
                else:
                    doc.add_paragraph(f"{key}: {value}")
            doc.add_paragraph('')
        
        # Risk analysis
        if risk_analysis:
            doc.add_heading('АНАЛИЗ РИСКОВ', 2)
            doc.add_paragraph(risk_analysis)
            doc.add_paragraph('')
        
        # Footer
        doc.add_paragraph('')
        doc.add_paragraph('Отчет подготовлен LEGALCHAIN AI')
        doc.add_paragraph('ПЕРЕПРОВЕРЬ все выводы перед использованием в суде!')
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def generate_court_filing(
        self,
        case_id: str,
        case_info: Dict[str, Any],
        key_facts: Dict[str, Any],
        timeline: List[Dict[str, Any]],
        discrepancies: List[Dict[str, Any]]
    ) -> BytesIO:
        """
        Generate Court Filing document (Word)
        
        Args:
            case_id: Case identifier
            case_info: Case information
            key_facts: Key facts dictionary
            timeline: Timeline events
            discrepancies: Discrepancies list
            
        Returns:
            BytesIO buffer with Word document
        """
        doc = Document()
        
        # Court header format
        doc.add_paragraph('В АРБИТРАЖНЫЙ СУД')
        if key_facts.get('court', {}).get('name'):
            doc.add_paragraph(key_facts['court']['name'])
        doc.add_paragraph('')
        
        # Parties
        if key_facts.get('parties'):
            doc.add_paragraph(f"Истец: {key_facts['parties'].get('plaintiff', 'Не указан')}")
            doc.add_paragraph(f"Ответчик: {key_facts['parties'].get('defendant', 'Не указан')}")
        doc.add_paragraph('')
        
        # Main content
        doc.add_heading('ОБОСНОВАНИЕ ПОЗИЦИИ', 2)
        doc.add_paragraph('На основании анализа документов дела:')
        doc.add_paragraph('')
        
        # Timeline
        if timeline:
            doc.add_heading('ХРОНОЛОГИЯ СОБЫТИЙ', 3)
            for event in timeline:
                doc.add_paragraph(f"{event.get('date', '')}: {event.get('description', '')}")
            doc.add_paragraph('')
        
        # Evidence
        if discrepancies:
            doc.add_heading('ДОКАЗАТЕЛЬСТВА', 3)
            for disc in discrepancies:
                doc.add_paragraph(f"- {disc.get('description', '')}")
                if disc.get('source_documents'):
                    doc.add_paragraph(f"  (Источники: {', '.join(disc['source_documents'])})")
            doc.add_paragraph('')
        
        # Footer
        doc.add_paragraph('')
        doc.add_paragraph('ВАЖНО: Этот документ сгенерирован автоматически.')
        doc.add_paragraph('ПЕРЕСМОТРИ перед подачей в суд!')
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def generate_contract_comparison(
        self,
        case_id: str,
        contracts_data: List[Dict[str, Any]]
    ) -> BytesIO:
        """
        Generate Contract Comparison report (Excel)
        
        Args:
            case_id: Case identifier
            contracts_data: List of contract data dictionaries
            
        Returns:
            BytesIO buffer with Excel file
        """
        wb = Workbook()
        ws = wb.active
        ws.title = "Сравнение контрактов"
        
        # Headers
        headers = ['Поставщик', 'Сумма', 'Оплата', 'Штраф', 'Гарантия', 'Статус']
        ws.append(headers)
        
        # Style headers
        header_fill = PatternFill(start_color="667eea", end_color="667eea", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center")
        
        # Add data
        for contract in contracts_data:
            row = [
                contract.get('supplier', ''),
                contract.get('amount', ''),
                contract.get('payment', ''),
                contract.get('penalty', ''),
                contract.get('warranty', ''),
                contract.get('status', '')
            ]
            ws.append(row)
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)
        return buffer
    
    def generate_pdf_report(
        self,
        case_id: str,
        summary: str,
        key_facts: Dict[str, Any]
    ) -> BytesIO:
        """
        Generate PDF report
        
        Args:
            case_id: Case identifier
            summary: Summary text
            key_facts: Key facts dictionary
            
        Returns:
            BytesIO buffer with PDF file
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        story = []
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#667eea'),
            alignment=1,  # Center
            spaceAfter=30
        )
        
        # Title
        story.append(Paragraph('LEGALCHAIN AI', title_style))
        story.append(Paragraph('EXECUTIVE SUMMARY', styles['Heading1']))
        story.append(Spacer(1, 12))
        
        # Case info
        story.append(Paragraph(f'Дело: {case_id}', styles['Normal']))
        story.append(Paragraph(f'Дата: {datetime.now().strftime("%d.%m.%Y")}', styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Summary
        story.append(Paragraph('СУТЬ ДЕЛА', styles['Heading2']))
        story.append(Paragraph(summary, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Key facts
        if key_facts:
            story.append(Paragraph('КЛЮЧЕВЫЕ ФАКТЫ', styles['Heading2']))
            if key_facts.get('parties'):
                story.append(Paragraph(f"Истец: {key_facts['parties'].get('plaintiff', 'Не указан')}", styles['Normal']))
                story.append(Paragraph(f"Ответчик: {key_facts['parties'].get('defendant', 'Не указан')}", styles['Normal']))
            story.append(Spacer(1, 20))
        
        # Footer
        story.append(Spacer(1, 30))
        story.append(Paragraph('Отчет подготовлен LEGALCHAIN AI', styles['Normal']))
        story.append(Paragraph('ПЕРЕПРОВЕРЬ все выводы перед использованием в суде!', styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer

